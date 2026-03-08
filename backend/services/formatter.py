"""
formatter.py — TemplateFormatter
==================================
IEEE A4 Formatter.  All formatting values read from ieee.json.

Bug fixes in this version (verified against test_output.docx forensics)
────────────────────────────────────────────────────────────────────────
BUG-7  _write_sections() SKIP set:
       Old: SKIP = {REFERENCES, BIBLIOGRAPHY, ACKNOWLEDGMENT, ACKNOWLEDGMENTS,
                    ACKNOWLEDGEMENTS} only.
            TITLE / AUTHORS / ABSTRACT / KEYWORDS were not in the set
            → written as "I. TITLE", "II. AUTHORS", "III. ABSTRACT", "IV. KEYWORDS"
            in the body column even after the header zone.
       Fix: expanded SKIP set covers all structural headings.

BUG-8  _write_acknowledgments():
       Old: ack_text = mapped_content["acknowledgments"] = "ACKNOWLEDGMENTS"
            (the mapper BUG-4 label string).  Because it is truthy the fallback
            section search was never reached → wrote "ACKNOWLEDGMENTS" as body.
       Fix: strip the acknowledgment label string before the truthy check; the
            section fallback is now reached for any non-body ack_text value.
            (With mapper BUG-4 also fixed this becomes a belt-and-suspenders guard.)

BUG-9  _write_tables() fully commented out:
       Old: entire method was commented out → tables produced only caption lines,
            no actual Word table objects.
       Fix: method restored and uncommented; includes XML border fallback (no
            add_style) when "Table Grid" is absent from the loaded template.
"""

import logging
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# All headings that must NEVER be written as Roman-numeral body sections
_SKIP_IN_BODY = {
    "TITLE", "AUTHORS", "AUTHOR",
    "ABSTRACT",
    "KEYWORDS", "KEYWORD", "INDEX TERMS",
    "REFERENCES", "BIBLIOGRAPHY",
    "ACKNOWLEDGMENT", "ACKNOWLEDGMENTS", "ACKNOWLEDGEMENTS",
}


class TemplateFormatter:
    """IEEE A4 Formatter with 4-Zone Layout Architecture."""

    def __init__(self, ieee_spec_path: str, template_path: str = None):
        with open(ieee_spec_path, "r", encoding="utf-8") as f:
            self._spec = json.load(f)

        if template_path and Path(template_path).exists():
            self.doc = Document(template_path)
            self._clear_body()
        else:
            self.doc = Document()
            logger.warning("No template loaded — using Normal style fallbacks.")

        self.stats = {
            "paragraphs_processed":    0,
            "paragraphs_corrected":    0,
            "paragraphs_skipped":      0,
            "properties_corrected":    0,
            "corrections_by_property": {},
        }

    # ── Helpers ────────────────────────────────────────────────────────────────

    def _clear_body(self) -> None:
        body = self.doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "sectPr":
                body.remove(child)

    # ── Public entry point ─────────────────────────────────────────────────────

    def format_document(self, mapped_content: dict) -> Document:
        if not mapped_content:
            raise ValueError("No mapped content provided.")

        self.mapped_content  = mapped_content
        self.paragraphs_data = mapped_content.get("paragraphs", [])

        # Build style-id index once so _find_para_data never does a text-prefix
        # match AFTER zone-writer methods have injected "Abstract—" etc. prefixes
        self._style_id_index: dict = {}
        for p in self.paragraphs_data:
            sid = p.get("parsed_style_id")
            if sid and sid not in self._style_id_index:
                self._style_id_index[sid] = p

        # ── ZONE 1: Full-width — title ────────────────────────────────────────
        self._write_title(mapped_content)
        self._insert_zone_break("zone1_title")

        # ── ZONE 2: Authors + Affiliations (Table-based layout) ──────────────
        self._write_authors_and_affiliations(mapped_content)
        self._insert_zone_break("zone2_affiliations", num_cols=1)

        # ── ZONE 3: Unified 2-column body (Abstract -> Keywords -> Sections) ──
        self._write_abstract(mapped_content)
        self._write_index_terms(mapped_content)

        # ── ZONE 4: Two-column body ───────────────────────────────────────────
        self._write_sections(mapped_content)
        self._write_tables(mapped_content)            # BUG-9 restored
        self._write_acknowledgments(mapped_content)
        self._write_references(mapped_content)

        self._apply_final_sectpr(mapped_content)
        return self.doc

    # ── Zone-break / sectPr ────────────────────────────────────────────────────

    def _insert_zone_break(self, zone_key: str, num_cols: int = None) -> None:
        zone = self._spec["zone_layout"][zone_key]
        if not self.doc.paragraphs:
            return
        last_para = self.doc.paragraphs[-1]
        pPr = last_para._p.get_or_add_pPr()
        for existing in pPr.findall(qn("w:sectPr")):
            pPr.remove(existing)
        sectPr = OxmlElement("w:sectPr")
        if "section_type" in zone and zone["section_type"]:
            type_el = OxmlElement("w:type")
            type_el.set(qn("w:val"), zone["section_type"])
            sectPr.append(type_el)
        pl = self._spec["page_layout"]
        pgSz = OxmlElement("w:pgSz")
        pgSz.set(qn("w:w"),    str(zone["page_w_twips"]))
        pgSz.set(qn("w:h"),    str(zone["page_h_twips"]))
        pgSz.set(qn("w:code"), str(pl.get("page_code", 9)))
        sectPr.append(pgSz)
        pgMar = OxmlElement("w:pgMar")
        pgMar.set(qn("w:top"),    str(int(zone["top_margin_pt"] * 20)))
        pgMar.set(qn("w:right"),  str(int(zone["right_pt"]       * 20)))
        pgMar.set(qn("w:bottom"), str(int(zone["bottom_pt"]      * 20)))
        pgMar.set(qn("w:left"),   str(int(zone["left_pt"]        * 20)))
        pgMar.set(qn("w:header"), str(int(zone["header_pt"]      * 20)))
        pgMar.set(qn("w:footer"), str(int(zone["footer_pt"]      * 20)))
        pgMar.set(qn("w:gutter"), str(pl.get("gutter_twips", 0)))
        sectPr.append(pgMar)
        
        cols_el = OxmlElement("w:cols")
        cols_val = num_cols if num_cols is not None else zone["cols"]
        cols_el.set(qn("w:num"),   str(cols_val))
        
        # Adjust gap based on number of columns
        if cols_val > 1:
            cols_el.set(qn("w:space"), str(zone["gap_twips"]))
        else:
            cols_el.set(qn("w:space"), "0")
            
        sectPr.append(cols_el)
        docGrid = OxmlElement("w:docGrid")
        docGrid.set(qn("w:linePitch"), str(pl.get("doc_grid_line_pitch", 360)))
        sectPr.append(docGrid)
        pPr.append(sectPr)
        
        # Force a pagination/section break in XML if needed
        # (Usually sectPr in pPr of last paragraph of zone is enough for Word)

    def _apply_final_sectpr(self, mapped_content: dict) -> None:
        zone = self._spec["zone_layout"]["zone3_body"]
        sp   = self.doc.sections[-1]._sectPr
        pl   = self._spec["page_layout"]
        for tag in ["w:cols", "w:pgSz", "w:pgMar", "w:type"]:
            for el in sp.findall(qn(tag)):
                sp.remove(el)

        if "section_type" in zone and zone["section_type"]:
            type_el = OxmlElement("w:type")
            type_el.set(qn("w:val"), zone["section_type"])
            sp.append(type_el)

        pgSz = OxmlElement("w:pgSz")
        pgSz.set(qn("w:w"),    str(zone["page_w_twips"]))
        pgSz.set(qn("w:h"),    str(zone["page_h_twips"]))
        pgSz.set(qn("w:code"), str(pl.get("page_code", 9)))
        sp.append(pgSz)
        pgMar = OxmlElement("w:pgMar")
        pgMar.set(qn("w:top"),    str(int(zone["top_margin_pt"] * 20)))
        pgMar.set(qn("w:right"),  str(int(zone["right_pt"]       * 20)))
        pgMar.set(qn("w:bottom"), str(int(zone["bottom_pt"]      * 20)))
        pgMar.set(qn("w:left"),   str(int(zone["left_pt"]        * 20)))
        pgMar.set(qn("w:header"), str(int(zone["header_pt"]      * 20)))
        pgMar.set(qn("w:footer"), str(int(zone["footer_pt"]      * 20)))
        pgMar.set(qn("w:gutter"), str(pl.get("gutter_twips", 0)))
        sp.append(pgMar)
        cols_el = OxmlElement("w:cols")
        cols_el.set(qn("w:num"),   str(zone["cols"]))
        cols_el.set(qn("w:space"), str(zone["gap_twips"]))
        sp.append(cols_el)
        docGrid = OxmlElement("w:docGrid")
        docGrid.set(qn("w:linePitch"), str(pl.get("doc_grid_line_pitch", 360)))
        sp.append(docGrid)

    # ── Zone writers ───────────────────────────────────────────────────────────

    def _write_title(self, mapped_content):
        title_text = mapped_content.get("title", "").strip()
        if not title_text:
            return

        cfg = self._spec["styles"].get("paper_title", {})
        font_name = cfg.get("font", "Times New Roman")
        font_size = cfg.get("size_pt", 24)
        
        p = self.doc.add_paragraph(title_text.upper()) # IEEE titles are UPPERCASE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(cfg.get("space_before_pt", 0))
        p.paragraph_format.space_after = Pt(cfg.get("space_after_pt", 6))

        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold      = cfg.get("bold", False)

    def _write_authors_and_affiliations(self, mapped: dict) -> None:
        """
        Groups each author name with their corresponding affiliation lines
        and writes them side-by-side using an invisible table.
        This provides robust horizontal alignment across all Word versions.
        """
        authors = mapped.get("authors", [])
        affs    = mapped.get("affiliations", [])
        if not authors:
            return

        a_cfg = self._spec["styles"].get("author", {})
        f_cfg = self._spec["styles"].get("affiliation", {})
        
        a_font = a_cfg.get("font", "Times New Roman")
        a_size = a_cfg.get("size_pt", 11)
        f_font = f_cfg.get("font", "Times New Roman")
        f_size = f_cfg.get("size_pt", 9)

        # Unicode superscript digits
        _SUP_RE = re.compile(r"([\u00b9\u00b2\u00b3\u2074-\u2079\u2070\u2071\u207f\u2020\u2021\u00a7*]+)$")

        # Create author table
        table = self.doc.add_table(rows=1, cols=len(authors))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Make it full width
        table.width = self.doc.sections[-1].page_width - self.doc.sections[-1].left_margin - self.doc.sections[-1].right_margin
        
        # Remove default borders
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
            
        look = OxmlElement("w:tblLook")
        look.set(qn("w:val"), "04A0") # Hidden borders
        tblPr.append(look)
        
        # Remove visual borders explicitly
        borders = OxmlElement("w:tblBorders")
        for b in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            el = OxmlElement(f"w:{b}")
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            borders.append(el)
        tblPr.append(borders)

        for i, author_raw in enumerate(authors):
            cell = table.cell(0, i)
            # Clear default paragraph
            cell.paragraphs[0].text = ""
            
            # 1. Split name from trailing superscript
            m_a = _SUP_RE.search(author_raw.strip())
            if m_a:
                name = author_raw[:m_a.start()].strip()
                a_sup = m_a.group(1)
            else:
                name = author_raw.strip()
                a_sup = None

            # 2. Write Author Name in cell
            p_name = cell.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_name.paragraph_format.space_before = Pt(a_cfg.get("space_before_pt", 12))
            p_name.paragraph_format.space_after  = Pt(a_cfg.get("space_after_pt", 2))
            
            r_name = p_name.add_run(name)
            r_name.font.name = a_font
            r_name.font.size = Pt(a_size)
            r_name.bold      = a_cfg.get("bold", False)

            # 3. Find and Write matching Affiliation parts in cell
            matching_aff = None
            if a_sup:
                for aff_str in affs:
                    if aff_str.startswith(a_sup):
                        matching_aff = aff_str[len(a_sup):].strip() # remove prefix correctly
                        break
            
            if matching_aff:
                parts = [p.strip() for p in matching_aff.split(", ")]
                for part in parts:
                    if not part: continue
                    p_aff = cell.add_paragraph()
                    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_aff.paragraph_format.space_before = Pt(0)
                    p_aff.paragraph_format.space_after  = Pt(0)
                    
                    r_aff = p_aff.add_run(part)
                    r_aff.font.name = f_font
                    r_aff.font.size = Pt(f_size)
                    r_aff.italic    = f_cfg.get("italic", True)
            
            # Remove the first dummy paragraph if cell.add_paragraph was used
            if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text:
                p = cell.paragraphs[0]._element
                p.getparent().remove(p)

    def _write_authors(self, mapped: dict) -> None:
        # Replaced by _write_authors_and_affiliations
        pass

    def _write_affiliations(self, mc: dict) -> None:
        # Replaced by _write_authors_and_affiliations
        pass

    def _write_abstract(self, mc: dict):
        text = mc.get("abstract", "").strip()
        if not text:
            return
        cfg = self._spec["styles"].get("abstract", {})
        font_name = cfg.get("font", "Times New Roman")
        font_size = cfg.get("size_pt", 9)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(cfg.get("space_before_pt", 0))
        p.paragraph_format.space_after  = Pt(cfg.get("space_after_pt", 10))
        p.paragraph_format.first_line_indent = Pt(cfg.get("first_line_indent_pt", 13.6))

        # "Abstract—" prefix is bold + italic at run level
        r_prefix = p.add_run("Abstract\u2014")
        r_prefix.font.name = font_name
        r_prefix.font.size = Pt(font_size)
        r_prefix.bold = True
        r_prefix.italic = True

        r_body = p.add_run(text)
        r_body.font.name = font_name
        r_body.font.size = Pt(font_size)
        r_body.bold = cfg.get("bold", False)
        r_body.italic = cfg.get("italic", False)

    def _write_index_terms(self, mc: dict) -> None:
        """
        FIX 4 — Keywords / Index Terms rendering.

        Old problems:
          • If mapper returned "Keywords—multi-UAV, MARL" the method prepended
            another "Keywords—" label → output: "Keywords—Keywords—multi-UAV…"
          • _para() wrote text into the Word style's implicit run; the explicit
            font/size was never set on that run → document default applied.
          • "Index Terms—" is the official IEEE label; using "Keywords—" is also
            acceptable but must be consistent with what the mapper strips.

        Changes:
          • Strip any existing Keywords/Index Terms prefix from the mapper text
            before writing so we never duplicate the label.
          • Write label ("Index Terms—") as one bold+italic run, body as plain run.
          • Apply font/size explicitly on both runs via the cfg values.
          • Justified alignment (IEEE style).
        """
        raw_text = mc.get("keywords", "").strip()
        if not raw_text:
            return

        # Strip any prefix the mapper may have left in the string (idempotent)
        raw_text = re.sub(
            r"^(index\s+terms?|keywords?)\s*[\u2014\u2013\-:]?\s*",
            "", raw_text, flags=re.IGNORECASE,
        ).strip()
        if not raw_text:
            return

        cfg       = self._spec["styles"].get("keywords", {})
        font_name = cfg.get("font", "Times New Roman")
        font_size = cfg.get("size_pt", 9)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(cfg.get("space_before_pt", 0))
        p.paragraph_format.space_after  = Pt(cfg.get("space_after_pt", 6))
        p.paragraph_format.first_line_indent = Pt(cfg.get("first_line_indent_pt", 13.7))

        # Label run — bold + italic per user request ("Keywords")
        r_label = p.add_run("Keywords\u2014")
        r_label.font.name = font_name
        r_label.font.size = Pt(font_size)
        r_label.bold      = True
        r_label.italic    = True

        # Body run — italic per IEEE spec
        r_body = p.add_run(raw_text)
        r_body.font.name = font_name
        r_body.font.size = Pt(font_size)
        r_body.italic    = cfg.get("italic", True)

    def _write_sections(self, mc: dict):
        """
        BUG-7: Skip set now covers TITLE / AUTHORS / ABSTRACT / KEYWORDS in
        addition to REFERENCES / ACKNOWLEDGMENTS.
        """
        section_count = 1
        for section in mc.get("sections", []):
            raw   = section.get("heading", "").strip()
            # Strip any Roman-numeral prefix that may have been added by a
            # previous formatter pass (idempotency guard)
            clean = re.sub(r"^[IVXLCDM]+\.\s*", "", raw, flags=re.IGNORECASE).strip().upper()
            if any(k in clean for k in _SKIP_IN_BODY):
                continue
            self._write_section_heading(section, section_count)
            self._write_content_recursive(section.get("content", []), level=1, parent_num=str(section_count))
            section_count += 1

    def _write_section_heading(self, section: dict, num: int) -> None:
        """
        FIX 7 — Section heading rendering.

        Old problems:
          • _para() used the word_style_name from ieee.json which may not exist
            in the loaded document/template → style exception silently swallowed →
            heading appeared as Normal style (no bold, wrong size).
          • No explicit font/size/bold set on the run → doc-default applied.
          • Heading text was built from section["heading"] which already contained
            a Roman numeral if the parser wrote "I. INTRODUCTION" → formatter
            added another → "I. I. INTRODUCTION".

        Changes:
          • Strip any existing Roman-numeral prefix from heading text before
            prepending the formatter's own numbering (idempotent).
          • Explicit run: Times New Roman, size from cfg (default 10 pt), bold.
          • Small-caps applied via XML vertAlign so it works without a named style.
          • Centered alignment (IEEE conference heading style).
          • Space before/after from cfg.
        """
        raw_heading = section.get("heading", "").strip()
        # Strip any pre-existing Roman numeral prefix added by parser or a
        # previous formatter pass ("I. INTRODUCTION" → "INTRODUCTION")
        clean_heading = re.sub(
            r"^(XIV|XIII|XII|XI|X{0,3}I{0,3}|IX|VIII|VII|VI|IV|V?I{1,3}|II|I)\.\s*",
            "", raw_heading, flags=re.IGNORECASE,
        ).strip().upper()

        heading_text = f"{self._roman(num)}. {clean_heading}"

        cfg       = self._spec["styles"].get("heading1", {})
        font_name = cfg.get("font", "Times New Roman")
        font_size = cfg.get("size_pt", 10)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before    = Pt(cfg.get("space_before_pt", 8))
        p.paragraph_format.space_after     = Pt(cfg.get("space_after_pt", 4))
        p.paragraph_format.keep_with_next  = cfg.get("keep_next", True)

        run = p.add_run(heading_text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold      = cfg.get("bold", True)

        # Apply small-caps via XML if ieee.json requests it
        if cfg.get("small_caps", False):
            rPr = run._r.get_or_add_rPr()
            sc  = OxmlElement("w:smallCaps")
            sc.set(qn("w:val"), "true")
            rPr.append(sc)

    def _write_content_recursive(
        self, content_list: list, level: int = 1, parent_num: str = ""
    ) -> None:
        """
        FIX 8 — Body paragraph and sub-heading rendering.

        Old problems:
          • Body paragraphs used _para(text, style_name) which relies on a named
            Word style that may not exist in the blank document → falls back to
            Normal → Calibri 11 pt, no first-line indent.
          • Sub-heading numbering produced "1.1 Clean Heading" but then also
            recursed into child content with level+1, causing grandchild sections
            to get numbering like "1.1.1" even for lettered sub-headings (A. B.).
          • Sub-heading clean regex r"^[\\d\\.\\sA-Z]+\\.\\s*" was too aggressive:
            it stripped the first word of headings like "Adaptive Control" →
            "daptive Control" (the leading "A." was eaten).

        Changes:
          • Body paragraphs: explicit Times New Roman, size_pt from cfg (10 pt),
            justified, first-line indent from cfg (default 14.4 pt ≈ 0.2 inch).
            space_before = 0, space_after = 2 pt for tight IEEE body spacing.
          • Sub-heading: preserve original heading text if it already has a
            numeric (3.1) or alpha (A.) prefix; only auto-number when the heading
            has no existing prefix.
          • Sub-heading run: bold, 10 pt, left-aligned (IEEE sub-heading style).
        """
        b_cfg     = self._spec["styles"].get("body_text", {})
        b_font    = b_cfg.get("font", "Times New Roman")
        b_size    = b_cfg.get("size_pt", 10)
        b_indent  = b_cfg.get("first_line_indent_pt", 14.4)

        h2_cfg    = self._spec["styles"].get("heading2", {})
        h2_font   = h2_cfg.get("font", "Times New Roman")
        h2_size   = h2_cfg.get("size_pt", 10)

        # Compiled patterns to detect an existing numeric/alpha prefix
        _HAS_NUM_PREFIX   = re.compile(r"^\d+\.\d+")
        _HAS_ALPHA_PREFIX = re.compile(r"^[A-Z]\.\s")

        sub_count = 1
        for item in content_list:
            if "heading" in item or "subheading" in item:
                raw_hd = (item.get("heading") or item.get("subheading", "")).strip()

                # Only auto-prefix when the heading has no existing numbering
                if _HAS_NUM_PREFIX.match(raw_hd) or _HAS_ALPHA_PREFIX.match(raw_hd):
                    heading_text = raw_hd  # already has "3.1 Foo" or "A. Foo"
                else:
                    if level == 1:
                        prefix = f"{parent_num}.{sub_count}" if parent_num else str(sub_count)
                    else:
                        prefix = self._alpha(sub_count)
                    heading_text = f"{prefix} {raw_hd}"

                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before     = Pt(h2_cfg.get("space_before_pt", 6))
                p.paragraph_format.space_after      = Pt(h2_cfg.get("space_after_pt", 3))
                p.paragraph_format.first_line_indent = Pt(h2_cfg.get("first_line_indent_pt", 0))
                p.paragraph_format.keep_with_next   = h2_cfg.get("keep_next", True)

                run = p.add_run(heading_text)
                run.font.name = h2_font
                run.font.size = Pt(h2_size)
                run.bold      = h2_cfg.get("bold", False)
                run.italic    = h2_cfg.get("italic", True)

                child_content = item.get("content", [])
                if child_content:
                    self._write_content_recursive(child_content, level + 1, heading_text.split()[0])
                sub_count += 1

            elif "text" in item:
                text = item["text"].strip()
                if not text:
                    continue

                role = item.get("role", "body")

                # Table/figure captions get their own alignment
                if role in {"table_caption", "figure_caption"}:
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after  = Pt(3)
                    run = p.add_run(text)
                    run.font.name = b_font
                    run.font.size = Pt(b_size - 2)  # captions 8 pt
                    run.bold      = True
                else:
                    # Standard body paragraph
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_before      = Pt(b_cfg.get("space_before_pt", 0))
                    p.paragraph_format.space_after       = Pt(b_cfg.get("space_after_pt", 6))
                    p.paragraph_format.first_line_indent = Pt(b_indent)

                    run = p.add_run(text)
                    run.font.name = b_font
                    run.font.size = Pt(b_size)

    def _write_acknowledgments(self, mc: dict):
        """
        BUG-8: Old code used ack_text = mc["acknowledgments"] directly.
        With mapper BUG-4, that value was "ACKNOWLEDGMENTS" (the heading label),
        which is truthy → the section fallback was never reached → the label
        was written as body text.

        Fix: strip known label strings before the truthy check; then try the
        section fallback if still empty.
        """
        ack_text = mc.get("acknowledgments", "").strip()

        # Strip bare heading label if mapper BUG-4 was not fixed on their end
        if ack_text.upper() in {"ACKNOWLEDGMENT", "ACKNOWLEDGMENTS", "ACKNOWLEDGEMENTS"}:
            ack_text = ""

        # Fallback: search sections list (handles both mapper-fixed and unfixed)
        if not ack_text:
            for section in mc.get("sections", []):
                h = section.get("heading", "").strip().upper()
                # Strip Roman-numeral prefix (idempotency)
                h_clean = re.sub(r"^[IVXLCDM]+\.\s*", "", h).strip()
                if any(k in h_clean for k in {
                    "ACKNOWLEDGMENT", "ACKNOWLEDGMENTS", "ACKNOWLEDGEMENTS"
                }):
                    ack_text = "\n".join(
                        item.get("text", "")
                        for item in section.get("content", [])
                        if "text" in item
                    ).strip()
                    break

        if not ack_text:
            return

        h_cfg = self._spec["styles"]["heading5"]
        self._para("ACKNOWLEDGMENT", h_cfg["word_style_name"])

        b_cfg = self._spec["styles"]["body_text"]
        self._para(ack_text, b_cfg["word_style_name"])

    def _write_tables(self, mapped_content: dict) -> None:
        """
        Robust IEEE table rendering.
        """
        tables = mapped_content.get("tables", [])
        if not tables:
            return

        for idx, table_data in enumerate(tables):
            # 1. Parse content into rows
            rows = table_data.get("rows", [])
            
            # Fallback to 'content' if 'rows' is missing (older mapper logic)
            if not rows:
                raw_content = table_data.get("content", [])
                if not raw_content:
                    continue
                for line in raw_content:
                    txt = line.strip()
                    if '\t' in txt:
                        cells = [c.strip() for c in txt.split('\t')]
                    else:
                        # Robust split: try 3+ spaces first, then 2+
                        cells = [p.strip() for p in re.split(r'\s{3,}', txt) if p.strip()]
                        if len(cells) < 2:
                            cells = [p.strip() for p in re.split(r'\s{2,}', txt) if p.strip()]
                    
                    if len(cells) >= 2:
                        # Unit merging consistent with parser/mapper
                        refined = []
                        for p in cells:
                            if refined and (p.startswith("(") or p.lower() in {"%", "(%)", "(ms)", "(kb)", "(mb)", "(ours)"}):
                                refined[-1] = f"{refined[-1]} {p}"
                            else:
                                refined.append(p)
                        rows.append(refined)

            if not rows:
                continue

            # 2. Determine grid size
            num_cols = max(len(r) for r in rows)
            # Pad short rows
            for r in rows:
                while len(r) < num_cols:
                    r.append("")

            # 3. Write Caption
            caption_text = table_data.get("caption", "")
            cleaned_caption = self._clean_caption(caption_text)
            
            p_cap = self.doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before   = Pt(12)
            p_cap.paragraph_format.space_after    = Pt(6)
            p_cap.paragraph_format.keep_with_next = True

            # Label (TABLE I.) - Bold
            run_label = p_cap.add_run(f"TABLE {self._roman(idx+1)}. ")
            run_label.font.name = "Times New Roman"
            run_label.font.size = Pt(8)
            run_label.bold = True

            # Title
            run_title = p_cap.add_run(cleaned_caption.upper())
            run_title.font.name = "Times New Roman"
            run_title.font.size = Pt(8)

            # 4. Create Table
            table_obj = self.doc.add_table(rows=len(rows), cols=num_cols)
            
            # IEEE Tables are centered
            tblPr = table_obj._tbl.find(qn("w:tblPr"))
            if tblPr is not None:
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "center")
                tblPr.append(jc)

            # 5. Fill Cells
            for r_idx, row_data in enumerate(rows):
                for c_idx, text in enumerate(row_data):
                    cell = table_obj.cell(r_idx, c_idx)
                    # Clear any default P
                    cell.text = "" 
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    run = p.add_run(str(text))
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8)
                    if r_idx == 0: # Header
                        run.bold = True

            # 6. Apply Borders (Crucial for visibility)
            self._apply_table_borders(table_obj)

    def _apply_table_borders(self, table) -> None:
        """XML border fallback — ensures consistent borders on all sides and inside."""
        tbl   = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
            
        # Clear existing borders if any
        existing_borders = tblPr.find(qn("w:tblBorders"))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
            
        borders = OxmlElement("w:tblBorders")
        for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "4") # 1/2 pt
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            borders.append(el)
        tblPr.append(borders)

    def _clean_caption(self, caption: str) -> str:
        cleaned = re.sub(
            r"^(table|tbl)[\s\.\-]*\w*[\s\.\:\-]+",
            "", caption.strip(), flags=re.IGNORECASE
        ).strip()
        return cleaned if cleaned else caption.strip()

    def _write_references(self, mc: dict) -> None:
        """
        FIX 6 — References rendering.

        Old problems:
          • Duplicate "[1] [1]" because the mapper already stored "[1] Author…"
            and the formatter prepended another "[1] " → "[1] [1] Author…".
          • re.sub(r"^\\[?\\d+\\]?[\\s\\.]*", …) is too greedy: it strips
            "1" from "10pt" in the first word when the reference starts without
            a bracket → wrong text.
          • No hanging indent: all lines of a multi-line reference wrap flush-left,
            which looks wrong compared to IEEE style (first line starts with [1],
            continuation lines indented ~18 pt).
          • Deduplication was done by the final refs list only — if the same
            paper appeared as "[3] Smith…" and "[3] Smith…" twice from the parser
            it was still included twice.

        Changes:
          • Strip leading bracket-number via a tighter regex: only removes the
            leading "[N]" or "N." token if it is the very first token and purely
            numeric (no false stripping of "10pt" etc.).
          • Deduplicate by normalised body text (not by the numbered prefix) so
            the same reference cannot appear twice regardless of what number it
            had in the source.
          • Hanging indent applied via paragraph_format: first_line_indent = -18pt,
            left indent = 18pt.  This is the standard Word approach; it works
            even without a named paragraph style.
          • REFERENCES heading written with heading1 style (falls back to heading5
            if heading1 not defined) for consistent small-caps / centered look.
        """
        refs = mc.get("references", [])
        if not refs:
            return

        # ── REFERENCES heading ─────────────────────────────────────────────
        h_cfg = self._spec["styles"].get(
            "heading1",
            self._spec["styles"].get("heading5", {})
        )
        h_p = self.doc.add_paragraph()
        h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_p.paragraph_format.space_before = Pt(6)
        h_p.paragraph_format.space_after  = Pt(6)
        h_run = h_p.add_run("REFERENCES")
        h_run.font.name = h_cfg.get("font", "Times New Roman")
        h_run.font.size = Pt(h_cfg.get("size_pt", 10))
        h_run.bold      = h_cfg.get("bold", True)

        # ── Regex to strip leading bracket number: [1], 1., [12], 12. ─────
        _LEAD_NUM = re.compile(r"^\[?\d+\]?[\.\s]+")

        r_cfg     = self._spec["styles"].get("references", {})
        font_name = r_cfg.get("font", "Times New Roman")
        font_size = r_cfg.get("size_pt", 8)

        # ── Deduplicate by normalised body text ────────────────────────────
        seen_bodies: set[str] = set()
        clean_refs: list[str] = []
        for ref in refs:
            ref = ref.strip()
            if not ref:
                continue
            body = _LEAD_NUM.sub("", ref).strip()
            body_key = re.sub(r"\s+", " ", body).lower()
            if body_key and body_key not in seen_bodies:
                seen_bodies.add(body_key)
                clean_refs.append(body)

        # ── Write each reference ───────────────────────────────────────────
        for i, body in enumerate(clean_refs, 1):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Hanging indent: left=18pt, first_line=-18pt
            pf = p.paragraph_format
            pf.left_indent        = Pt(18)
            pf.first_line_indent  = Pt(-18)
            pf.space_before       = Pt(0)
            pf.space_after        = Pt(2)

            run = p.add_run(f"[{i}] {body}")
            run.font.name = font_name
            run.font.size = Pt(font_size)

    # ── Para helpers ───────────────────────────────────────────────────────────

    def _para(self, text: str, style_name: str):
        p = self.doc.add_paragraph(text)
        try:
            p.style = style_name
        except Exception:
            pass
        return p

    def _find_para_data(
        self,
        parsed_style_id: str = None,
        role: str            = None,
        text_prefix: str     = None,
    ):
        """Style-id index lookup first, then linear fallback."""
        if parsed_style_id and parsed_style_id in self._style_id_index:
            return self._style_id_index[parsed_style_id]
        for p in self.paragraphs_data:
            if role and p.get("role") != role:
                continue
            if text_prefix and not p.get("text", "").startswith(text_prefix[:50]):
                continue
            return p
        return None

    def _apply_violations(self, paragraph, para_data):
        self.stats["paragraphs_processed"] += 1
        if not para_data or para_data.get("is_valid", True):
            self.stats["paragraphs_skipped"] += 1
            return

        self.stats["paragraphs_corrected"] += 1
        violations = para_data.get("violations", {})
        spec_props = para_data.get("spec_properties", {})

        if not paragraph.runs and paragraph.text:
            orig = paragraph.text.strip()
            p_el = paragraph._p
            for child in list(p_el):
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag in ("r", "t"):
                    p_el.remove(child)
            if orig:
                paragraph.add_run(orig)

        handlers = {
            "font_name":            self._h_font_name,
            "size_pt":              self._h_size_pt,
            "bold":                 self._h_bold,
            "italic":               self._h_italic,
            "small_caps":           self._h_small_caps,
            "alignment":            self._h_alignment,
            "space_before_pt":      self._h_space_before,
            "space_after_pt":       self._h_space_after,
            "first_line_indent_pt": self._h_first_line_indent,
            "line_spacing_pt":      self._h_line_spacing,
            "line_spacing_rule":    self._h_line_spacing_rule,
        }
        for prop_key in violations:
            handler = handlers.get(prop_key)
            if handler:
                try:
                    handler(paragraph, spec_props[prop_key])
                    self.stats["properties_corrected"] += 1
                    self.stats["corrections_by_property"][prop_key] = (
                        self.stats["corrections_by_property"].get(prop_key, 0) + 1
                    )
                except Exception as e:
                    logger.error(f"Failed to apply {prop_key}: {e}")

    # ── Property handlers ──────────────────────────────────────────────────────

    def _h_font_name(self, p, val):
        for r in p.runs: r.font.name = val

    def _h_size_pt(self, p, val):
        for r in p.runs: r.font.size = Pt(val)

    def _h_bold(self, p, val):
        for r in p.runs: r.bold = val

    def _h_italic(self, p, val):
        for r in p.runs: r.italic = val

    def _h_small_caps(self, p, val):
        for r in p.runs:
            rPr = r._r.get_or_add_rPr()
            for sc in rPr.findall(qn("w:smallCaps")): rPr.remove(sc)
            if val:
                sc_el = OxmlElement("w:smallCaps")
                sc_el.set(qn("w:val"), "true")
                rPr.append(sc_el)

    def _h_alignment(self, p, val):
        m = {
            "center":  WD_ALIGN_PARAGRAPH.CENTER,
            "left":    WD_ALIGN_PARAGRAPH.LEFT,
            "right":   WD_ALIGN_PARAGRAPH.RIGHT,
            "both":    WD_ALIGN_PARAGRAPH.JUSTIFY,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "start":   WD_ALIGN_PARAGRAPH.LEFT,
        }
        p.alignment = m.get(val.lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)

    def _h_space_before(self, p, val):      p.paragraph_format.space_before       = Pt(val)
    def _h_space_after(self, p, val):       p.paragraph_format.space_after        = Pt(val)
    def _h_first_line_indent(self, p, val): p.paragraph_format.first_line_indent  = Pt(val)
    def _h_line_spacing(self, p, val):      p.paragraph_format.line_spacing       = Pt(val)

    def _h_line_spacing_rule(self, p, val):
        m = {
            "auto":     WD_LINE_SPACING.MULTIPLE,
            "exact":    WD_LINE_SPACING.EXACTLY,
            "multiple": WD_LINE_SPACING.MULTIPLE,
            "at_least": WD_LINE_SPACING.AT_LEAST,
        }
        p.paragraph_format.line_spacing_rule = m.get(
            val.lower(), WD_LINE_SPACING.MULTIPLE
        )

    # ── Utilities ──────────────────────────────────────────────────────────────

    def _roman(self, n: int) -> str:
        vals = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
        syms = ["M","CM","D","CD","C","XC","L","XL","X","IX","V","IV","I"]
        r, i = "", 0
        while n > 0:
            for _ in range(n // vals[i]):
                r += syms[i]
                n -= vals[i]
            i += 1
        return r

    def _alpha(self, n: int) -> str:
        return chr(64 + n) if 1 <= n <= 26 else str(n)

    def save_document(self, output_path: str) -> str:
        self.doc.save(output_path)
        return str(output_path)

    def get_stats(self) -> dict:
        return self.stats


# ── Local test ─────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    mapper_output = r"debug_mapped_fixed.json"
    output_path   = r"test_output.docx"
    ieee_path     = r"D:\DATA SCIENCE AND ML\Project\ManuscriptMagic.AI\backend\templates\ieee.json"
    template_path = r"D:\DATA SCIENCE AND ML\Project\ManuscriptMagic.AI\template\conference-template-a4.docx"

    with open(mapper_output, "r", encoding="utf-8") as f:
        data = json.load(f)

    fmt = TemplateFormatter(ieee_spec_path=ieee_path, template_path=template_path)
    fmt.format_document(data)
    fmt.save_document(output_path)

    print("Formatting Complete!")
    s = fmt.get_stats()
    print(f"Paragraphs corrected : {s['paragraphs_corrected']}")
    print(f"Paragraphs skipped   : {s['paragraphs_skipped']}")
    print(f"Properties corrected : {s['properties_corrected']}")
"""
parser.py — DocumentParser
===========================
Extracts structured content from a user's raw manuscript (.docx).

Handles ALL real-world document formats:
  - IEEE style (Roman numerals: I. Introduction)
  - Inline label + dash (Abstract— or Keywords—)
  - ALL-CAPS prefix (ABSTRACT Large Language..., TITLE Advancements...)
  - Standalone uppercase headings (INTRODUCTION, RELATED WORK)
  - Numeric sub-sections (3.1 Methodology, 4.2 Baselines)
  - Plain body text with no formatting metadata

Requirements:
    pip install python-docx
"""

import re
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ─────────────────────────────────────────────
#  Compiled Regex Patterns
# ─────────────────────────────────────────────

# ── Roman numeral section headings ──
# FIX: old regex was anchored wrong — couldn't match "I." alone because
#      group(1) consumed nothing and group(2) needed at least one char.
# NEW: explicit alternation covers I–XIII and common IEEE section numbers.
#      Allows optional trailing ALL-CAPS title: "I. INTRODUCTION"
_ROMAN_RE = re.compile(
    r"^(XIV|XIII|XII|XI|X{1,3}|IX|VIII|VII|VI|IV|V|V?I{1,3}|II|I)\."
    r"(\s+[A-Z][A-Z\s\-]{0,60})?$",
    re.IGNORECASE,
)

# Sub-heading: single capital letter + dot + space, short text
# Strict: only A–Z, not arbitrary letters starting sentences
_ALPHA_HEAD_RE = re.compile(r"^[A-Z]\.\s")

# Reference patterns
_REF_BRACKET_RE  = re.compile(r"^\[\d+\]")          # [1] Author ...
_REF_NUMERIC_RE  = re.compile(r"^\d+\.\s")           # 1. Author ...

# Arabic numbered main section: 1. Introduction, 2. Methodology, 4. Results
_ARABIC_MAIN_RE = re.compile(
    r"^\d+\.\s+([A-Z][A-Za-z\s&\-]{4,80})$",
    re.IGNORECASE
)

# Numeric sub-section full match: 2.1 Early Approaches, 3.2.1 Dataset
_NUMERIC_SUB_FULL_RE = re.compile(
    r"^\d+\.\d+(\.\d+)?\s+([A-Z][A-Za-z\s\-]{4,100})$",
    re.IGNORECASE
)

# Numeric sub-section (prefix): "3.1 Title", "4.2 Baselines", "3.1.2 Detail"
_NUMERIC_SUB_RE  = re.compile(r"^\d+\.\d+(\.\d+)?\s+\S")

# ── Abstract patterns ──
# Covers: "Abstract—...", "Abstract: ...", "Abstract ...", "ABSTRACT ..."
# FIX: added bold-marker strip for "**Abstract**—" docx bold runs
_ABSTRACT_INLINE_RE  = re.compile(r"^abstract[\s\u2014\u2013\-:]\s*", re.IGNORECASE)
_ABSTRACT_CAPS_SPLIT = re.compile(r"^ABSTRACT\s+", re.MULTILINE)  # "ABSTRACT body"

# ── Keywords patterns ──
# Covers: "Keywords—...", "Keywords: ...", "Index Terms—...", "KEYWORDS ..."
# FIX: added em-dash U+2014 and en-dash U+2013 variants explicitly
_KEYWORDS_INLINE_RE  = re.compile(
    r"^(keywords?|index terms?)[\s\u2014\u2013\-:]\s*", re.IGNORECASE
)

# ── ALL-CAPS inline label detection ──
# Matches paragraphs starting with an ALL-CAPS label word(s) then a space.
# "TITLE Advancements..." / "AUTHORS Harshit..." / "INTRODUCTION The rapid..."
# Label must be 2–30 chars, all uppercase letters/spaces, followed by content.
_CAPS_LABEL_RE = re.compile(
    r"^([A-Z][A-Z\s]{1,28}?)\s{1,3}(?=[A-Za-z0-9])"
)

# ── Table / Figure caption ──
# FIX: extended to match "TABLE I.", "TABLE 1.", "TABLE I:" and bold variants.
_TABLE_CAPTION_RE = re.compile(
    r"^(TABLE|TBL\.?)\s+([IVX]+|\d+)[.\-:\s]", re.IGNORECASE
)
_FIGURE_CAPTION_RE = re.compile(
    r"^(FIG\.?|FIGURE)\s+([IVX]+|\d+)[.\-:\s]", re.IGNORECASE
)

# ── Author name line ──
# FIX: detects lines like "Rahul Sharma¹" or "John A. Smith², Jane Doe³"
# Pattern: 2–5 words in Title Case optionally followed by superscript digits/symbols.
_AUTHOR_NAME_RE = re.compile(
    r"^([A-Z][a-z]+(?:\s+[A-Z]\.?)?(?:\s+[A-Z][a-z]+)+)"  # First [Middle] Last
    r"[\u00b9\u00b2\u00b3\u00b4\u2070-\u2079\u207a-\u207f,\s]*$"  # optional superscripts
)

# Superscript affiliation number at line start: "¹Department...", "²IIT Delhi"
_AFF_SUPERSCRIPT_START_RE = re.compile(
    r"^[\u00b9\u00b2\u00b3\u00b4\u2070-\u2079\u207a-\u207f]"
)

# Email line
_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")

# Set of valid ALL-CAPS inline labels we recognise
_CAPS_SECTION_LABELS = {
    "TITLE", "AUTHORS", "AUTHOR",
    "ABSTRACT",
    "KEYWORDS", "KEYWORD", "INDEX TERMS",
    "INTRODUCTION",
    "RELATED WORK", "BACKGROUND", "LITERATURE REVIEW",
    "METHODOLOGY", "METHOD", "METHODS", "APPROACH",
    "PROPOSED METHOD", "PROPOSED METHODOLOGY",
    "EXPERIMENTAL SETUP", "EXPERIMENTS", "EXPERIMENTAL RESULTS",
    "RESULTS", "RESULTS AND DISCUSSION",
    "DISCUSSION", "ANALYSIS", "EVALUATION",
    "CONCLUSION", "CONCLUSIONS", "FUTURE WORK",
    "CONCLUSION AND FUTURE WORK",
    "ACKNOWLEDGMENT", "ACKNOWLEDGMENTS", "ACKNOWLEDGEMENTS",
    "REFERENCES", "BIBLIOGRAPHY",
}

# Standalone uppercase section names (paragraph = just the heading text)
_KNOWN_HEADINGS_UPPER = {l.upper() for l in {
    "abstract", "introduction", "related work", "background",
    "literature review", "methodology", "methods", "approach",
    "proposed method", "system design", "architecture",
    "experimental results", "experiments", "results",
    "results and discussion", "discussion", "analysis",
    "evaluation", "conclusion", "conclusions",
    "future work", "conclusion and future work",
    "acknowledgment", "acknowledgements", "acknowledgments",
    "references", "bibliography",
}}

# Lowercase version for case-insensitive matching
_KNOWN_HEADINGS = {h.lower() for h in _KNOWN_HEADINGS_UPPER}

# ── Affiliation signals ──
# FIX: added "india", "china", "usa", "email" as common affiliation tokens
#      added "professor", "researcher", "engineer" job-title signals
_AFFILIATION_WORD_SIGNALS = {
    "university", "institute", "department", "dept", "college",
    "laboratory", "lab", "centre", "center", "school", "faculty",
    "iit", "nit", "iisc", "stanford", "caltech", "oxford",
    "hospital", "clinic", "corporation", "corp", "ltd",
    # Country / city names that almost always appear only in affiliations
    "india", "usa", "uk", "china", "germany", "france", "japan",
    "canada", "australia", "brazil", "korea",
    # Job-title signals
    "professor", "researcher", "engineer", "scientist", "lecturer",
    "email",
}

_AFFILIATION_SUBSTR_SIGNALS = {
    "@", ".edu", ".ac.", ".org", ".gov",
    "\u00b9", "\u00b2", "\u00b3", "\u2020", "\u2021", "\u00a7",
}


# ─────────────────────────────────────────────
#  DocumentParser
# ─────────────────────────────────────────────

class DocumentParser:
    """
    Parses a user's raw manuscript DOCX and returns structured content
    suitable for the ContentClassifier and TemplateFormatter.

    Usage:
        parser = DocumentParser("path/to/paper.docx")
        content = parser.extract_all()
    """

    def __init__(self, filepath: str | Path):
        self.filepath = Path(filepath)
        if not self.filepath.is_file():
            raise FileNotFoundError(f"Document not found: {self.filepath}")
        
        self.is_json = self.filepath.suffix.lower() == ".json"
        
        if not self.is_json:
            try:
                self.doc = Document(self.filepath)
            except Exception as e:
                print(f"❌ Failed to load DOCX (Zip Error?): {e}")
                raise
        else:
            self.doc = None
            print(f"ℹ️ Input is JSON, bypassing DOCX parsing: {self.filepath}")

    def extract_all(self) -> dict[str, Any]:
        """Runs the full extraction pipeline."""
        if self.is_json:
            try:
                with open(self.filepath, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                print(f"❌ Failed to load JSON input: {e}")
                raise

        raw_paragraphs = self._extract_paragraphs()
        
        # Build caption candidate list in order for Change 7
        self._caption_candidates = [
            p["text"] for p in raw_paragraphs 
            if p["role"] == "table_caption" or re.search(r"\btable\s+(i{1,3}|iv|v?i{0,3}|\d+)\b", p["text"], re.I)
        ]

        tables  = self._extract_tables(raw_paragraphs)
        images  = self._extract_images()

        title        = self._find_title(raw_paragraphs)
        authors      = self._find_authors(raw_paragraphs)
        affiliations = self._find_affiliations(raw_paragraphs)
        abstract     = self._find_abstract(raw_paragraphs)
        keywords     = self._find_keywords(raw_paragraphs)
        sections     = self._build_sections(raw_paragraphs)
        references   = self._find_references(raw_paragraphs)

        # New: Extract text-based tables from paragraphs if native tables are missing
        if not tables:
            tables = self._extract_tables_from_text(raw_paragraphs)

        return {
            "source_file"   : self.filepath.name,
            "title"         : title,
            "authors"       : authors,
            "affiliations"  : affiliations,
            "abstract"      : abstract,
            "keywords"      : keywords,
            "sections"      : sections,
            "references"    : references,
            "tables"        : tables,
            "images"        : images,
            "raw_paragraphs": raw_paragraphs,
            "metadata"      : {
                "paragraph_count" : len(raw_paragraphs),
                "table_count"     : len(tables),
                "image_count"     : len(images),
                "section_count"   : len(sections),
                "reference_count" : len(references),
                "word_count"      : sum(len(p["text"].split()) for p in raw_paragraphs),
            },
        }

    def to_json(self, pretty: bool = True, save_path=None) -> str:
        data     = self.extract_all()
        json_str = json.dumps(data, indent=2 if pretty else None, ensure_ascii=False)
        if save_path:
            Path(save_path).write_text(json_str, encoding="utf-8")
        return json_str

    # ──────────────────────────────────────────────────────
    #  Step 1: Paragraph Extraction
    # ──────────────────────────────────────────────────────

    def _extract_paragraphs(self) -> list[dict[str, Any]]:
        result = []
        index  = 0
        
        _ARTIFACT_BLOCKLIST = {
            "END OF PAPER", "END OF DOCUMENT", "- - -", "---",
            "* * *", "***"
        }

        # Map roles to style IDs for Problem 5
        role_to_style_id = {
            "title": "paper_title",
            "author": "author",
            "affiliation": "body_text",
            "abstract": "abstract",
            "keywords": "keywords",
            "section_heading": "heading1",
            "sub_heading": "heading2",
            "body": "body_text",
            "reference": "references",
            "abstract_heading": "abstract",
            "keywords_heading": "keywords",
        }

        for para in self.doc.paragraphs:
            raw_text = para.text.strip()
            if not raw_text:
                continue

            text = self._clean_text(raw_text)
            
            # Bug 4 Fix: Blocklist filter
            if text.upper() in _ARTIFACT_BLOCKLIST:
                continue

            # SAFE STYLE ACCESS: para.style can be None in some documents
            word_style_name = para.style.name if para.style is not None else "Normal"
            resolved_props = self._read_ooxml_properties(para)
            run_overrides = self._read_run_overrides(para, resolved_props)

            # Bug 1 & 2 Fix: split now returns a list of dicts
            split_nodes = self._split_inline_headings(text)

            for node in split_nodes:
                part_text   = node["text"]
                forced_role = node["forced_role"]

                # Detect ALL-CAPS inline label (legacy but kept for caps_label/content metadata)
                caps_label, caps_content = self._split_caps_label(part_text)

                if forced_role:
                    role = forced_role
                else:
                    prev_role = result[-1]["role"] if result else None
                    role = self._classify_paragraph_role(
                        part_text, caps_label, 
                        resolved_props["size_pt"], 
                        resolved_props["bold"], 
                        resolved_props["alignment"], 
                        word_style_name,
                        prev_role
                    )

                parsed_style_id = role_to_style_id.get(role, "body_text")

                entry = {
                    "index"               : index,
                    "text"                : part_text,
                    "word_style_name"     : word_style_name,
                    "parsed_style_id"     : parsed_style_id,
                    "role"                : role,
                    "ooxml_properties"    : resolved_props,
                    "run_level_overrides" : run_overrides,
                }

                if caps_label:
                    entry["caps_label"]   = caps_label
                    entry["caps_content"] = caps_content

                result.append(entry)
                index += 1

        # Bug 3 Fix: Post-classification override pass for references
        in_references = False
        for entry in result:
            text_upper = entry["text"].strip().upper()
            if entry["role"] == "section_heading" and text_upper in {"REFERENCES", "BIBLIOGRAPHY"}:
                in_references = True
                continue
            
            if in_references:
                if entry["role"] in {"body", "section_heading"}:
                    # New section started? 
                    if entry["role"] == "section_heading" and text_upper not in {"REFERENCES", "BIBLIOGRAPHY"}:
                        in_references = False
                    else:
                        entry["role"] = "reference"
                        entry["parsed_style_id"] = "references"

        return result

    def _read_ooxml_properties(self, para: Paragraph) -> dict:
        """
        Walks the full inheritance chain: paragraph level -> style -> parent styles -> defaults.
        Implements Change 1.
        """
        # IEEE Body Text Defaults as ultimate fallback
        defaults = {
            "font_name": "Times New Roman",
            "size_pt": 10.0,
            "bold": False,
            "italic": False,
            "small_caps": False,
            "alignment": "both",
            "space_before_pt": 0.0,
            "space_after_pt": 6.0,
            "first_line_indent_pt": 14.4,
            "line_spacing_pt": 11.4,
            "line_spacing_rule": "auto"
        }

        def get_prop(obj, attr, is_font=True):
            try:
                if is_font:
                    val = getattr(obj.font, attr)
                else:
                    val = getattr(obj.paragraph_format, attr)
                return val
            except Exception:
                return None

        # 1. Alignment Mapping
        align_map = {
            WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
        }
        
        resolved = {}
        
        # ─── FONT PROPERTIES ───
        # Check first run as proxy for paragraph-level default if para-style is missing it
        font_name = None
        size_pt = None
        bold = None
        italic = None
        small_caps = None

        # Chain walk for fonts
        chain = []
        if para.style:
            chain.append(para.style)
            curr = para.style
            try:
                while curr and hasattr(curr, 'base_style') and curr.base_style:
                    chain.append(curr.base_style)
                    curr = curr.base_style
            except Exception:
                pass # Broken style chain in Word XML
        
        # Check runs for overrides that might be "paragraph wide"
        if para.runs:
            first_run = para.runs[0]
            if first_run.font:
                font_name = first_run.font.name
                if first_run.font.size:
                    size_pt = round(first_run.font.size.pt, 1)
                bold = first_run.bold
                italic = first_run.italic
                small_caps = first_run.font.small_caps

        # Follow style chain if run properties are None
        for style in chain:
            if not style or not hasattr(style, 'font') or style.font is None:
                continue
            if font_name is None: font_name = style.font.name
            if size_pt is None and style.font.size: size_pt = round(style.font.size.pt, 1)
            if bold is None: bold = style.font.bold
            if italic is None: italic = style.font.italic
            if small_caps is None: small_caps = style.font.small_caps

        # Final Doc Defaults for font
        if font_name is None or size_pt is None:
            try:
                doc_defaults = self.doc.element.find(qn('w:docDefaults'))
                if doc_defaults is not None:
                    rPr = doc_defaults.find(qn('w:rPrDefault')).find(qn('w:rPr'))
                    if font_name is None:
                        rFont = rPr.find(qn('w:rFonts'))
                        if rFont is not None: font_name = rFont.get(qn('w:ascii'))
                    if size_pt is None:
                        sz = rPr.find(qn('w:sz'))
                        if sz is not None: size_pt = int(sz.get(qn('w:val'))) / 2.0
            except Exception: pass

        resolved["font_name"] = font_name or defaults["font_name"]
        resolved["size_pt"] = size_pt or defaults["size_pt"]
        resolved["bold"] = bold if bold is not None else defaults["bold"]
        resolved["italic"] = italic if italic is not None else defaults["italic"]
        resolved["small_caps"] = small_caps if small_caps is not None else defaults["small_caps"]

        # ─── PARAGRAPH PROPERTIES ───
        alignment = para.alignment
        space_before = para.paragraph_format.space_before
        space_after = para.paragraph_format.space_after
        first_line_indent = para.paragraph_format.first_line_indent
        line_spacing = para.paragraph_format.line_spacing
        line_spacing_rule_raw = para.paragraph_format.line_spacing_rule

        for style in chain:
            if alignment is None: alignment = style.paragraph_format.alignment
            if space_before is None: space_before = style.paragraph_format.space_before
            if space_after is None: space_after = style.paragraph_format.space_after
            if first_line_indent is None: first_line_indent = style.paragraph_format.first_line_indent
            if line_spacing is None: line_spacing = style.paragraph_format.line_spacing
            if line_spacing_rule_raw is None: line_spacing_rule_raw = style.paragraph_format.line_spacing_rule

        # Alignment mapping
        resolved["alignment"] = align_map.get(alignment, defaults["alignment"])
        
        # Spacing conversions
        resolved["space_before_pt"] = round(space_before.pt, 1) if space_before else defaults["space_before_pt"]
        resolved["space_after_pt"] = round(space_after.pt, 1) if space_after else defaults["space_after_pt"]
        resolved["first_line_indent_pt"] = round(first_line_indent.pt, 1) if first_line_indent else defaults["first_line_indent_pt"]
        resolved["line_spacing_pt"] = round(line_spacing * resolved["size_pt"], 1) if line_spacing else defaults["line_spacing_pt"]
        
        # line_spacing_rule: "exact" if EXACTLY, "auto" if AT_LEAST/proportional (multi-line), "multiple" otherwise
        # WD_LINE_SPACING enum: 0=SINGLE, 1=1.5, 2=DOUBLE, 3=AT_LEAST, 4=EXACTLY, 5=MULTIPLE
        from docx.enum.text import WD_LINE_SPACING
        if line_spacing_rule_raw == WD_LINE_SPACING.EXACTLY:
            resolved["line_spacing_rule"] = "exact"
        elif line_spacing_rule_raw in {WD_LINE_SPACING.AT_LEAST, WD_LINE_SPACING.SINGLE, WD_LINE_SPACING.ONE_POINT_FIVE, WD_LINE_SPACING.DOUBLE}:
            resolved["line_spacing_rule"] = "auto"
        else:
            resolved["line_spacing_rule"] = "multiple"

        return resolved

    def _read_run_overrides(self, para: Paragraph, resolved_props: dict) -> dict:
        """
        Detects runs with different font_name or size_pt than the paragraph base.
        Implements Change 2.
        """
        runs_out = []
        has_mixed = False
        
        for i, run in enumerate(para.runs):
            if run.font is None:
                continue
                
            run_name = run.font.name
            run_size = round(run.font.size.pt, 1) if run.font.size else None
            
            is_diff = False
            if run_name and run_name != resolved_props["font_name"]:
                is_diff = True
            if run_size and run_size != resolved_props["size_pt"]:
                is_diff = True
            
            if is_diff:
                has_mixed = True
            
            # We always check for bold/italic diffs too if they are explicitly set on run
            if run.bold is not None and run.bold != resolved_props["bold"]:
                is_diff = True
            if run.italic is not None and run.italic != resolved_props["italic"]:
                is_diff = True

            if is_diff:
                runs_out.append({
                    "run_index": i,
                    "text": run.text,
                    "font_name": run_name or resolved_props["font_name"],
                    "size_pt": run_size or resolved_props["size_pt"],
                    "bold": run.bold if run.bold is not None else resolved_props["bold"],
                    "italic": run.italic if run.italic is not None else resolved_props["italic"],
                })

        return {
            "has_mixed_formatting": has_mixed,
            "runs": runs_out if has_mixed else []
        }

    def _split_inline_headings(self, text: str) -> list[dict[str, Any]]:
        """
        Splits fused section heading, sub-heading, and body text.
        Implements Bug 1 & 2 Fix.
        """
        label_to_role = {
            "TITLE": "title",
            "AUTHORS": "author", "AUTHOR": "author",
            "ABSTRACT": "abstract_heading",
            "KEYWORDS": "keywords_heading", "KEYWORD": "keywords_heading",
            "INDEX TERMS": "keywords_heading",
            "REFERENCES": "section_heading", "BIBLIOGRAPHY": "section_heading",
            "ACKNOWLEDGMENT": "section_heading", "ACKNOWLEDGMENTS": "section_heading", "ACKNOWLEDGEMENTS": "section_heading",
        }

        nodes = []
        remainder = text

        # Step 1: Check if text starts with a known caps section label.
        m = _CAPS_LABEL_RE.match(remainder)
        if m:
            label = m.group(1).strip()
            if label in _CAPS_SECTION_LABELS:
                role = label_to_role.get(label, "section_heading")
                nodes.append({"text": label, "forced_role": role})
                remainder = remainder[m.end():].strip()
        
        if not remainder:
            return nodes if nodes else [{"text": text, "forced_role": None}]

        # Step 2: From the remainder, check for numeric sub-heading
        m_num = re.match(r"^(\d+\.\d+(?:\.\d+)?)\s+(\S.*)$", remainder)
        if m_num:
            prefix = m_num.group(1)
            content_after_num = m_num.group(2).strip()

            # Try to find first sentence end after short heading-like prefix
            m_split = re.match(r"^([A-Z][A-Za-z0-9\s\.\-]{4,80})\.\s+(.*)", content_after_num)
            
            if m_split:
                potential_head = m_split.group(1).strip()
                rest = m_split.group(2).strip()
                # Only split if rest starts with lowercase or clear new sentence
                if rest and (rest[0].islower() or re.match(r"^[A-Z][a-z].*\.\s", rest)):
                    nodes.append({"text": f"{prefix} {potential_head}.", "forced_role": "sub_heading"})
                    nodes.append({"text": rest, "forced_role": None})
                    remainder = ""
            
            if remainder:
                # Check if it's just a short title
                words = content_after_num.split()
                if len(words) <= 12:
                    nodes.append({"text": f"{prefix} {content_after_num}", "forced_role": "sub_heading"})
                    remainder = ""
                else:
                    nodes.append({"text": f"{prefix} {content_after_num}", "forced_role": None})
                    remainder = ""

        if remainder:
            nodes.append({"text": remainder, "forced_role": None})

        return nodes if nodes else [{"text": text, "forced_role": None}]

    # ──────────────────────────────────────────────────────
    #  Step 2: ALL-CAPS label splitter
    # ──────────────────────────────────────────────────────

    def _split_caps_label(self, text: str) -> tuple[str | None, str | None]:
        """
        Detects ALL-CAPS inline label prefix.

        "TITLE Advancements in RAG..." → ("TITLE", "Advancements in RAG...")
        "ABSTRACT Large Language..."   → ("ABSTRACT", "Large Language...")
        "INTRODUCTION The rapid..."    → ("INTRODUCTION", "The rapid...")
        "Cancer remains..."            → (None, None)

        Returns (label, content) or (None, None) if no label detected.
        """
        m = _CAPS_LABEL_RE.match(text)
        if not m:
            return None, None

        label = m.group(1).strip()
        if label not in _CAPS_SECTION_LABELS:
            return None, None

        content = text[m.end():].strip()
        return label, content

    # ──────────────────────────────────────────────────────
    #  Step 3: Role Classification
    # ──────────────────────────────────────────────────────

    def _classify_paragraph_role(
        self,
        text: str,
        caps_label: str | None,
        font_size: float | None,
        is_bold: bool,
        alignment: str,
        style_name: str,
        prev_role: str | None = None,
    ) -> str:
        """
        Priority-ordered rule chain. Handles all real-world formats.

        Rule order (highest priority first):
          0  Context carry-over (author/affiliation block)
          1  Word style name (ground truth)
          2  ALL-CAPS inline label
          3  Inline Abstract (Abstract—...)
          4  Inline Keywords (Keywords—...)
          5  Standalone abstract / keywords heading
          6  Reference patterns  ([1], 1.)
          7  Roman numeral section heading  (I. INTRODUCTION)
          8  Standalone known heading (exact match)
          9  Standalone ALL-CAPS known heading
         10  Numeric sub-section  (3.1 Title)
         11  Alpha sub-heading    (A. Title)
         12  Table / figure caption
         13  Author name heuristic (superscript OR title-case name pattern)
         14  Affiliation heuristic
         15  Title heuristic (large, bold, centered)
         16  Default → body
        """
        style_lower = style_name.lower()
        text_lower  = text.lower().strip()
        text_stripped = text.strip()

        # ── Rule 0: Context carry-over for author / affiliation block ──
        # Once we are inside an author-block, keep classifying adjacent lines
        # as author or affiliation until a structural break occurs.
        if prev_role in {"author", "affiliation"}:
            # Hard breaks — these cannot be author/affiliation regardless
            _HARD_BREAKS = {
                "abstract", "abstract_heading", "keywords", "keywords_heading",
                "section_heading", "sub_heading", "reference",
            }
            # Check hard-break patterns first (roman numeral, bracket ref, etc.)
            if (
                _ROMAN_RE.match(text_stripped)
                or _REF_BRACKET_RE.match(text_stripped)
                or _ABSTRACT_INLINE_RE.match(text_lower)
                or _KEYWORDS_INLINE_RE.match(text_lower)
                or text_lower in _KNOWN_HEADINGS
            ):
                pass  # Fall through to normal rules below
            else:
                # Superscript-prefixed line → affiliation
                if _AFF_SUPERSCRIPT_START_RE.match(text_stripped):
                    return "affiliation"
                # Email line → affiliation
                if _EMAIL_RE.search(text_stripped):
                    return "affiliation"
                # Affiliation signals → affiliation
                if self._looks_like_affiliation(text):
                    return "affiliation"
                # Title-Case name pattern (2–5 words) → author
                if _AUTHOR_NAME_RE.match(text_stripped):
                    return "author"

        # ── Rule 1: Word styles (ground truth when present) ──
        if "papertitle" in style_lower or style_lower == "title":
            return "title"
        if style_lower == "author":
            return "author"
        if style_lower == "abstract":
            return "abstract"
        if style_lower == "keywords":
            return "keywords"
        if style_lower == "references":
            return "reference"
        if "heading 1" in style_lower or "heading1" in style_lower:
            return "section_heading"
        if "heading 2" in style_lower or "heading2" in style_lower:
            return "sub_heading"
        if "heading 3" in style_lower or "heading3" in style_lower:
            return "sub_heading"

        # ── Rule 2: ALL-CAPS inline label (legacy new-format support) ──
        if caps_label:
            if caps_label in {"TITLE"}:
                return "title"
            if caps_label in {"AUTHORS", "AUTHOR"}:
                return "author"
            if caps_label in {"ABSTRACT"}:
                return "abstract"
            if caps_label in {"KEYWORDS", "KEYWORD", "INDEX TERMS"}:
                return "keywords"
            if caps_label in {
                "INTRODUCTION", "RELATED WORK", "BACKGROUND", "LITERATURE REVIEW",
                "METHODOLOGY", "METHOD", "METHODS", "APPROACH",
                "PROPOSED METHOD", "PROPOSED METHODOLOGY",
                "EXPERIMENTAL SETUP", "EXPERIMENTAL RESULTS", "EXPERIMENTS",
                "RESULTS", "RESULTS AND DISCUSSION",
                "DISCUSSION", "ANALYSIS", "EVALUATION",
                "CONCLUSION", "CONCLUSIONS", "FUTURE WORK",
                "CONCLUSION AND FUTURE WORK",
                "ACKNOWLEDGMENT", "ACKNOWLEDGMENTS", "ACKNOWLEDGEMENTS",
            }:
                return "section_heading"

        # ── Rule 3: Inline Abstract with dash / colon / space ──
        # Handles: "Abstract—text", "Abstract: text", "ABSTRACT text"
        if _ABSTRACT_INLINE_RE.match(text_lower):
            return "abstract"

        # ── Rule 4: Inline Keywords ──
        # Handles: "Keywords—k1, k2", "Index Terms—k1, k2"
        if _KEYWORDS_INLINE_RE.match(text_lower):
            return "keywords"

        # ── Rule 5: Standalone abstract / keywords heading ──
        if text_lower in {"abstract", "abstract\u2014", "abstract:"}:
            return "abstract_heading"
        if re.match(r"^(keywords?|index terms?)\s*[\u2014\u2013\-:]?\s*$", text_lower):
            return "keywords_heading"

        # ── Rule 6: Reference patterns ──
        if "reference" in style_lower or "bibliography" in style_lower:
            return "reference"
        if _REF_BRACKET_RE.match(text_stripped):
            return "reference"
        if _REF_NUMERIC_RE.match(text_stripped):
            body_part = _REF_NUMERIC_RE.sub('', text_stripped).strip()
            if (len(text_stripped) > 35
                and not _ARABIC_MAIN_RE.match(text_stripped)
                and not body_part.upper().startswith(('INTRODUC', 'METHOD', 'RESULT', 'EXPERIM', 'CONCLUS', 'RELATED'))):
                return "reference"

        # ── Rule 7: Roman numeral section heading ──
        # FIX: new _ROMAN_RE correctly matches "I.", "II.", "I. INTRODUCTION" etc.
        if _ROMAN_RE.match(text_stripped):
            return "section_heading"

        # ── Rule 7b: Arabic numbered main section ───────────────────────────────
        if _ARABIC_MAIN_RE.match(text_stripped):
            heading_part = _ARABIC_MAIN_RE.match(text_stripped).group(1).strip()
            heading_upper = heading_part.upper()
            
            if (heading_upper.startswith(('INTRODUCTION', 'RELATED', 'METHOD', 'EXPERIMENT', 'RESULT', 'DISCUSS', 'CONCLUS'))
                or (len(heading_part.split()) <= 10
                    and heading_part[0].isupper()
                    and not heading_part.endswith(('.', ':', ';', '—', '-')))):
                return "section_heading"

        # ── Rule 8: Standalone known heading (exact match, any case) ──
        if text_lower in _KNOWN_HEADINGS:
            return "section_heading"

        # ── Rule 9: Standalone ALL-CAPS known heading without inline content ──
        if text_stripped.upper() == text_stripped and text_stripped in _KNOWN_HEADINGS_UPPER:
            return "section_heading"

        # ── Rule 10: Numeric sub-section heading ("3.1 Title", "4.2 Baselines") ──
        if _NUMERIC_SUB_FULL_RE.match(text_stripped):
            heading_part = _NUMERIC_SUB_FULL_RE.match(text_stripped).group(2).strip()
            if len(heading_part.split()) <= 12:
                return "sub_heading"
        if _NUMERIC_SUB_RE.match(text_stripped) and len(text_stripped.split()) <= 12:
            return "sub_heading"

        # ── Rule 11: Alpha sub-heading ("A. Feature Extraction") ──
        if _ALPHA_HEAD_RE.match(text_stripped) and len(text_stripped.split()) <= 8:
            return "sub_heading"

        # ── Rule 12: Table / Figure caption ──
        # FIX: refined to reject "mentions" (e.g. "Table I shows...")
        if _TABLE_CAPTION_RE.match(text_stripped):
            # A real caption is usually short (< 25 words) and doesn't end with a verb phrase like "shows"
            words = text_stripped.split()
            if len(words) < 25 and not any(w in text_lower for w in ["shows", "illustrates", "depicts"]):
                return "table_caption"
        if _FIGURE_CAPTION_RE.match(text_stripped):
            words = text_stripped.split()
            if len(words) < 25:
                return "figure_caption"

        # ── Rule 13: Author name heuristic ──
        # FIX: two distinct sub-checks ordered safest-first.
        #
        # 13a. Superscript-prefixed line → definitely affiliation, not author.
        if _AFF_SUPERSCRIPT_START_RE.match(text_stripped):
            return "affiliation"
        #
        # 13b. Email → affiliation.
        if _EMAIL_RE.search(text_stripped):
            return "affiliation"
        #
        # 13c. Title-Case name-only line with optional trailing superscripts.
        #      e.g. "Rahul Sharma¹"  "John A. Smith²"
        #      Guard: must be ≤ 7 words and not contain affiliation signals.
        if (
            _AUTHOR_NAME_RE.match(text_stripped)
            and len(text_stripped.split()) <= 7
            and not self._looks_like_affiliation(text)
        ):
            return "author"

        # ── Rule 14: Affiliation heuristic ──
        if self._looks_like_affiliation(text):
            return "affiliation"

        # ── Rule 15: Title heuristic (large bold centered text) ──
        if (
            font_size is not None and font_size >= 14
            and alignment == "center"
            and is_bold
            and len(text_stripped.split()) <= 20
        ):
            return "title"

        # ── Rule 13-legacy: Centered, short, no affiliation signals → author ──
        # Kept as a low-priority fallback only; real authors caught by Rule 13c above.
        if alignment == "center" and not self._looks_like_affiliation(text):
            if font_size is None or font_size <= 13:
                words = text_stripped.split()
                # Extra guard: reject single-word lines that are likely headings
                if 2 <= len(words) <= 6:
                    return "author"

        return "body"

    # ──────────────────────────────────────────────────────
    #  Step 4: Targeted Extractors
    # ──────────────────────────────────────────────────────

    def _find_title(self, paragraphs: list[dict]) -> str | None:
        for i, p in enumerate(paragraphs):
            if p["role"] == "title":
                text = p.get("caps_content") or p["text"]
                if text.strip().upper() in {"TITLE", "AUTHORS"}:
                    if i + 1 < len(paragraphs):
                        return paragraphs[i+1]["text"].strip()
                return text.strip()
        return None

    def _find_authors(self, paragraphs: list[dict]) -> list[str]:
        authors = []
        for p in paragraphs:
            if p["role"] == "author":
                text = p.get("caps_content") or p["text"]
                if text.strip().upper() in {"AUTHORS", "AUTHOR"}:
                    continue
                name, _ = self._split_author_affiliation(text)
                if name and name not in authors:
                    authors.append(name)
        return authors

    def _find_affiliations(self, paragraphs: list[dict]) -> list[str]:
        affs = []
        for p in paragraphs:
            if p["role"] == "affiliation" or any(word in p["text"].lower() for word in ["university", "institute", "department", "iit", "nit", "@", ".edu", "professor"]):
                affs.append(p["text"].strip())
        return affs

    def _split_author_affiliation(self, text: str) -> tuple[str, str | None]:
        """
        Heuristic: author name is the leading 1-4 words in Title Case; 
        everything after is the affiliation. (Change 6)
        """
        words = text.split()
        if not words:
            return text, None
            
        name_tokens = []
        aff_tokens = []
        
        for i, word in enumerate(words):
            # Author name words are Title Case and not signals
            is_title_case = word[0].isupper() and (len(word) == 1 or word[1:].islower() or word[1:].isalpha())
            is_signal = word.lower() in _AFFILIATION_WORD_SIGNALS
            
            if len(name_tokens) < 4 and is_title_case and not is_signal:
                name_tokens.append(word)
            else:
                aff_tokens = words[i:]
                break
                
        name = " ".join(name_tokens)
        aff = " ".join(aff_tokens) if aff_tokens else None
        return name, aff

    def _find_abstract(self, paragraphs: list[dict]) -> str | None:
        """
        Handles:
        1. Inline em-dash:  "Abstract—Full text..."
        2. Space separator: "ABSTRACT Full text..."
        3. Standalone heading + following body paragraphs
        4. Style-based:     style="Abstract"
        """
        abstract_parts = []
        in_abstract    = False

        for p in paragraphs:
            role = p["role"]
            text = p["text"]

            if role == "abstract":
                # Strip any prefix variant
                content = p.get("caps_content") or text
                cleaned = _ABSTRACT_INLINE_RE.sub("", content).strip()
                if not cleaned:
                    cleaned = content.strip()
                if cleaned:
                    return cleaned

            if role == "abstract_heading":
                in_abstract = True
                continue

            if in_abstract:
                if role == "body":
                    abstract_parts.append(text)
                else:
                    break

        return " ".join(abstract_parts) if abstract_parts else None

    def _find_keywords(self, paragraphs: list[dict]) -> str | None:
        """
        Handles:
        1. Inline em-dash:  "Keywords—deep learning, CNN"
        2. Space separator: "KEYWORDS deep learning, CNN"
        3. Standalone heading + next body paragraph
        """
        in_keywords = False

        for p in paragraphs:
            role = p["role"]
            text = p["text"]

            if role == "keywords":
                content = p.get("caps_content") or text
                cleaned = _KEYWORDS_INLINE_RE.sub("", content).strip()
                if not cleaned:
                    cleaned = content.strip()
                return cleaned if cleaned else None

            if role == "keywords_heading":
                in_keywords = True
                continue

            if in_keywords:
                if role == "body":
                    return text
                else:
                    break

        return None

    def _find_references(self, paragraphs: list[dict]) -> list[str]:
        """
        Collects references. Handles [1] brackets, numbered lists,
        author-year format (Author, Year), and body paragraphs
        inside the References section.
        """
        refs          = []
        in_ref_section = False

        for p in paragraphs:
            role = p["role"]
            text = p["text"]

            if role == "reference":
                refs.append(text)
                continue

            # Enter reference section
            if role == "section_heading" and text.upper() in {"REFERENCES", "BIBLIOGRAPHY"}:
                in_ref_section = True
                continue

            # Body paragraphs inside reference section = references
            if in_ref_section and role == "body":
                refs.append(text)

            # Exit on next non-reference heading
            if (role == "section_heading"
                    and text.upper() not in {"REFERENCES", "BIBLIOGRAPHY"}):
                in_ref_section = False

        # De-duplicate preserving order
        seen = set()
        unique = []
        for r in refs:
            if r not in seen:
                seen.add(r)
                unique.append(r)
        return unique

    def _build_sections(self, paragraphs: list[dict]) -> list[dict]:
        """
        Groups body content under section and sub-section headings.
        Heading must be caps_label, not caps_content. (Change 5)
        """
        SKIP = {
            "title", "author", "affiliation",
            "abstract", "abstract_heading",
            "keywords", "keywords_heading",
            "reference",
        }

        sections           = []
        current_section    = None
        current_subsection = None

        for p in paragraphs:
            role = p["role"]
            if role in SKIP or p.get("is_table_component"):
                continue

            if role == "section_heading":
                # Flush
                if current_subsection and current_section:
                    current_section["content"].append(current_subsection)
                    current_subsection = None
                if current_section:
                    sections.append(current_section)

                # Heading is caps_label if present (Problem 2 / Change 5)
                heading_text = p.get("caps_label") or p["text"]

                current_section = {
                    "heading"     : heading_text,
                    "heading_role": "section_heading",
                    "content"     : [],
                }

                # caps_content becomes the first body content item (Change 5)
                if p.get("caps_content"):
                    current_section["content"].append({
                        "text"     : p["caps_content"],
                        "role"     : "body",
                        "is_bold"  : False,
                        "is_italic": False,
                    })

            elif role == "sub_heading":
                if current_subsection and current_section:
                    current_section["content"].append(current_subsection)

                current_subsection = {
                    "heading"     : p["text"],
                    "heading_role": "sub_heading",
                    "content"     : [],
                }

            else:
                # Add body, figure_caption, table_caption etc.
                entry = {
                    "text"     : p["text"],
                    "role"     : role,
                    "is_bold"  : p["ooxml_properties"]["bold"],
                    "is_italic": p["ooxml_properties"]["italic"],
                }
                if current_subsection is not None:
                    current_subsection["content"].append(entry)
                elif current_section is not None:
                    current_section["content"].append(entry)

        # Flush remaining
        if current_subsection and current_section:
            current_section["content"].append(current_subsection)
        if current_section:
            sections.append(current_section)

        return sections

    # ──────────────────────────────────────────────────────
    #  Table + Image Extraction
    # ──────────────────────────────────────────────────────

    def _extract_tables(self, paragraphs: list[dict]) -> list[dict]:
        tables = []
        for idx, table in enumerate(self.doc.tables):
            rows = []
            for row in table.rows:
                row_data   = []
                seen_cells = set()
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    if cell_id in seen_cells:
                        continue
                    seen_cells.add(cell_id)
                    cell_text = " ".join(
                        p.text.strip() for p in cell.paragraphs if p.text.strip()
                    )
                    row_data.append(cell_text)
                if any(row_data):
                    rows.append(row_data)

            tables.append({
                "table_index"    : idx,
                "row_count"      : len(rows),
                "col_count"      : len(rows[0]) if rows else 0,
                "has_header_row" : self._detect_header_row(rows),
                "rows"           : rows,
                "caption"        : self._find_table_caption(idx, paragraphs),
            })

        return tables

    def _extract_tables_from_text(self, paragraphs: list[dict]) -> list[dict]:
        """
        Fallback for documents where tables are represented as plain text.
        Uses a numeric-ratio heuristic to detect data rows.
        """
        extracted_tables = []
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            text = p["text"].strip()
            
            if p["role"] == "table_caption":
                caption = text
                table_rows = []
                
                # Collect following rows that look like data
                j = i + 1
                consecutive_non_numeric = 0
                
                while j < len(paragraphs):
                    next_item = paragraphs[j]
                    next_text = next_item["text"].strip()
                    next_role = next_item.get("role", "body")
                    
                    if not next_text:
                        j += 1
                        continue
                        
                    # Stop if we hit a new structural element (but ignore "body")
                    if next_role not in {"body", "table_caption"}:
                        break

                    # Smarter row detection
                    # If we are in the first 2 rows, we allow non-numeric (headers)
                    # as long as they have at least 2 words.
                    cells = self._get_table_row_cells(next_text)
                    is_row = False
                    
                    if cells:
                        is_row = True
                    elif len(table_rows) < 2 and len(next_text.split()) >= 3:
                        # Fallback for single-spaced header rows
                        cells = next_text.split()
                        is_row = True

                    if is_row:
                        table_rows.append(cells)
                        j += 1
                        # If row is mostly numeric, reset the counter
                        if cells and sum(1 for w in cells if self._is_numeric(w)) / len(cells) > 0.3:
                            consecutive_non_numeric = 0
                        else:
                            consecutive_non_numeric += 1
                            
                        # If we hit too many non-numeric labels in a row, it's not a table
                        if consecutive_non_numeric > 3:
                            break
                    else:
                        break
                
                # Filter out "tables" that are just 1 paragraph of text
                if table_rows and (len(table_rows) > 1 or any(self._is_numeric(w) for c in table_rows for w in c)):
                    extracted_tables.append({
                        "table_index": len(extracted_tables),
                        "row_count": len(table_rows),
                        "col_count": max(len(r) for r in table_rows) if table_rows else 0,
                        "has_header_row": True,
                        "rows": table_rows,
                        "caption": caption
                    })
                    for idx in range(i, j):
                        paragraphs[idx]["is_table_component"] = True
                    i = j
                    continue
            i += 1
        return extracted_tables

    def _get_table_row_cells(self, text: str) -> list[str] | None:
        """
        Heuristic to split a string into table cells.
        Returns list of cells if it looks like a table row, else None.
        Matches IEEE/ACM multi-column text formatting artifacts.
        """
        # 1. Try splitting by Tab OR 3+ spaces first (most reliable for columns)
        parts = [p.strip() for p in re.split(r"\t|\s{3,}", text) if p.strip()]
        
        # 2. If no wide gaps, try 2+ spaces
        if len(parts) < 2:
            parts = [p.strip() for p in re.split(r"\t|\s{2,}", text) if p.strip()]
            
        if len(parts) < 2:
            return None
            
        # 3. Post-process: merge units and fragments that were over-split
        # e.g. ["Accuracy", "(%)"] -> ["Accuracy (%)"]
        # e.g. ["Latency", "(ms)"] -> ["Latency (ms)"]
        refined = []
        for p in parts:
            p_strip = p.strip()
            # Heuristic for unit markers or small trailing fragments
            is_unit_fragment = (
                p_strip.startswith("(") or 
                p_strip.endswith(")") or
                p_strip.lower() in {"%", "(%)", "(ms)", "(kb)", "(mb)", "ms", "kb", "mb", "(ours)"} or
                (len(p_strip) <= 4 and any(c in p_strip for c in "%()[]"))
            )
            
            if refined and is_unit_fragment:
                refined[-1] = f"{refined[-1]} {p_strip}"
            else:
                refined.append(p_strip)
                
        return refined if len(refined) >= 2 else None

    def _is_numeric(self, text: str) -> bool:
        """Checks if a string represents a number, percentage, or scientific value."""
        # Strip common units % ( ) [ ] ,
        clean = re.sub(r"[%()\-\[\],]", "", text)
        try:
            float(clean)
            return True
        except ValueError:
            return False

    def _detect_header_row(self, rows: list[list[str]]) -> bool:
        """
        A row is a header if all cells are non-numeric strings and the next row 
        contains at least one numeric. (Problem 7 / Change 8)
        """
        if len(rows) < 2:
            return False
            
        first_row = rows[0]
        second_row = rows[1]
        
        # Check first row has no numeric values
        for cell in first_row:
            if not cell: return False
            try:
                float(cell.replace(",", ""))
                return False # Is numeric
            except ValueError:
                pass
                
        # Check second row has at least one numeric value
        has_numeric = False
        for cell in second_row:
            try:
                float(cell.replace(",", ""))
                has_numeric = True
                break
            except ValueError:
                pass
                
        return has_numeric

    def _find_table_caption(self, table_idx: int, paragraphs: list[dict]) -> str | None:
        """
        Match captions to tables by order of appearance. (Problem 6 / Change 7)
        """
        if not hasattr(self, "_caption_candidates") or table_idx >= len(self._caption_candidates):
            return None
            
        raw_cap = self._caption_candidates[table_idx]
        
        # Strip leading section heading text (Problem 6)
        # Use a simple regex to find "TABLE [I|V|X|0-9]" and take from there
        m = re.search(r"\b(table\s+[i\d\.]+)\b", raw_cap, re.I)
        if m:
            return raw_cap[m.start():].strip()
            
        return raw_cap.strip()

    def _extract_images(self) -> list[dict]:
        images = []
        for idx, shape in enumerate(self.doc.inline_shapes):
            try:
                w = round(shape.width  / 914400, 3)
                h = round(shape.height / 914400, 3)
            except Exception:
                w = h = None
            images.append({
                "image_index"  : idx,
                "width_inches" : w,
                "height_inches": h,
                "shape_type"   : str(shape.shape_type),
            })
        return images

    # ──────────────────────────────────────────────────────
    #  Formatting helpers
    # ──────────────────────────────────────────────────────

    def _clean_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def _get_font_size(self, para: Paragraph) -> float | None:
        for run in para.runs:
            if run.font and run.font.size is not None:
                return round(run.font.size.pt, 1)
        try:
            if para.style and para.style.font and para.style.font.size is not None:
                return round(para.style.font.size.pt, 1)
        except Exception:
            pass
        return None

    def _get_font_name(self, para: Paragraph) -> str | None:
        for run in para.runs:
            if run.font and run.font.name:
                return run.font.name
        try:
            if para.style and para.style.font and para.style.font.name:
                return para.style.font.name
        except Exception:
            pass
        return None

    def _any_run_bold(self, para: Paragraph) -> bool:
        return any(run.bold for run in para.runs)

    def _any_run_italic(self, para: Paragraph) -> bool:
        return any(run.italic for run in para.runs)

    def _get_alignment(self, para: Paragraph) -> str:
        # Legacy method - _read_ooxml_properties handles this now
        return self._read_ooxml_properties(para)["alignment"]

    def _looks_like_affiliation(self, text: str) -> bool:
        """
        Returns True if the text looks like an institutional affiliation.

        FIX: raised word-count ceiling from 12 → 20 (real affiliations can be long:
             "Department of Computer Science, IIT Delhi, New Delhi, India").
        FIX: superscript-prefix check moved here so it is reachable from Rule 0.
        FIX: added email regex check and country/job-title signals (see module-level sets).
        """
        # Superscript prefix (¹²³ etc.) = affiliation line number
        if text and _AFF_SUPERSCRIPT_START_RE.match(text):
            return True

        # Email address → affiliation
        if _EMAIL_RE.search(text):
            return True

        # Reject very long paragraphs — they are body text, not affiliations
        if len(text.split()) > 20:
            return False

        text_lower = text.lower()

        # Substring signals (email domain, unicode superscripts)
        for signal in _AFFILIATION_SUBSTR_SIGNALS:
            if signal in text:
                return True

        # Word-boundary signals
        words = re.findall(r'\b\w+\b', text_lower)
        for word in words:
            if word in _AFFILIATION_WORD_SIGNALS:
                return True

        return False


if __name__ == "__main__":
    # ──────────────────────────────────────────────────────
    #  HARDCODED CONFIGURATION
    # ──────────────────────────────────────────────────────
    import os
    # Default to the conference template if no specific file is set
    INPUT_FILE_PATH = r"D:\DATA SCIENCE AND ML\Project\ManuscriptMagic.AI\template\test.docx"
    
    # This automatically creates a .json file with the same name in the same folder
    OUTPUT_JSON_PATH = Path(INPUT_FILE_PATH).with_suffix(".json")

    # 1. Check if file exists
    if not os.path.exists(INPUT_FILE_PATH):
        print(f"❌ Error: Could not find the file at:\n{INPUT_FILE_PATH}")
        print("\nPlease update the 'INPUT_FILE_PATH' variable in the script.")
    else:
        try:
            print(f"🚀 Starting Parser...")
            print(f"📄 Input: {INPUT_FILE_PATH}")

            # 2. Initialize and Run Parser
            doc_parser = DocumentParser(INPUT_FILE_PATH)
            extracted_content = doc_parser.extract_all()

            # 3. Save to JSON
            with open(OUTPUT_JSON_PATH, "w", encoding="utf-8") as f:
                json.dump(extracted_content, f, indent=2, ensure_ascii=False)
            
            print(f"✅ Success! Extracted content saved to:")
            print(f"📂 {OUTPUT_JSON_PATH}")
            
            # Optional: Preview the title
            if extracted_content.get("title"):
                print(f"\nExtracted Title: {extracted_content['title']}")

        except Exception as e:
            print(f"❌ An unexpected error occurred: {e}")